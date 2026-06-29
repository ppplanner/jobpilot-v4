"""
简历 API
- 上传简历文件 → 提取文本
- AI 诊断简历
- 获取/保存 简历版本
"""
import io
import sys
import os
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List

from core.database import get_db


def _llm_key() -> str:
    """读取 LLM API Key，优先从 settings 单例取，兜底到 os.environ。
    解决 settings 保存后单例未更新导致 resume 接口仍报「未配置」的问题。"""
    try:
        from core.config import settings
        return settings.llm_api_key or os.environ.get("LLM_API_KEY", "")
    except Exception:
        return os.environ.get("LLM_API_KEY", "")

# resume_analyzer 函数直接在本文件中定义（原jobpilot/utils已删除）
def extract_text_from_file(content: bytes, filename: str) -> str:
    """从文件字节提取纯文本（在upload端点内联实现，此处为兼容性保留）"""
    return content.decode("utf-8", errors="ignore")

def build_diagnosis_prompt(resume_text: str, target_role: str = "产品经理") -> str:
    return f"""你是一位资深HR和{target_role}求职顾问。请对以下简历进行专业诊断：

简历内容：
{resume_text}

请从以下维度诊断（输出结构化分析）：
1. 整体评分（1-10分）及理由
2. 亮点（至少3条）
3. 问题与不足（至少3条）
4. 针对{target_role}岗位的改进建议（至少5条具体建议）
5. 一句话总结"""

def build_rewrite_prompt(resume_text: str, jd_text: str, mode: str) -> str:
    """构建简历改写 prompt"""
    jd_section = f"\n\n【目标岗位 JD】\n{jd_text.strip()}" if jd_text.strip() else ""

    if mode == "light":
        instruction = """你是一位经验丰富的求职顾问，正在帮助候选人优化简历每段经历的呈现方式。

【改写原则 - 轻度优化】
核心要求：保留原有事实和数据，让语言更有力，结构更清晰，让HR扫一眼就抓住重点。

【每段经历的结构要求（最重要）】
每段经历的 rewritten 必须包含：
1. 总领句（第一行）：结果 / 交付物先行，行动在后
   HR筛简历时扫第一句话，「完成了什么」比「负责了什么」更有说服力
   格式参考：「[完成了什么 / 交付了什么]，[覆盖范围 / 独立程度]」
   示例（好）：「独立完成B端审批模块从0到1落地，覆盖财务/采购/行政3个业务部门」
   示例（差）：「负责B端审批模块产品设计，推动项目从0到1上线」← 动词先行，结果靠后
2. 2-3条具体要点（以「· 」开头）：
   第1条加一句背景（短语即可），说明你在解决什么问题——让读者知道事情的难度，避免「换个名字也能贴」的空洞描述
   示例：「· 针对旧系统操作繁琐导致供应商投诉频发，主导用户调研（访谈15人），定位核心流程卡点」
   第2-3条：核心行动 + 具体方法/工具/交付物，动词开头，20-40字

【rewritten 字段格式】
总领句\n· 要点1（背景+行动）\n· 要点2\n· 要点3（可选）

【改写要求】
✅ 可以做的：
- 将模糊动词换成精准行动动词（"参与"→"主导"、"负责"→"设计并推动"）
- 把最重要的信息放到总领句，让HR一眼看到核心价值
- 根据 JD 关键词，在原有表达中适当突出匹配点

❌ 绝对不可以做的：
- 不得添加原文没有的数字、百分比或规模数据
- 不得改变时间、公司名、岗位名等基本事实
- 不得使用以下词语（AI腔识别清单）：
  · 空洞强度词：深入、深度、充分、全面、系统性、多维度、有力、大力、有效（单独修饰时）
  · 企业黑话：赋能、闭环、打通、沉淀、对齐、拉通、助力、推动落地、抓手、底层逻辑
  · 自我评价：积极参与、深度参与、具有较强的、热爱、执行力强
  · 虚假结果：取得良好成效、获得一致好评、产生积极影响
  · AI句式：不得使用「通过…实现了…」「基于…推动了…」「不断优化…」等套路句式

【note 字段格式】
一句话说明改写侧重点（如"突出了需求分析能力，与JD匹配"）；若无JD则留空。

【输出格式（严格JSON）】
{
  "internships": [{"company": "公司名", "position": "岗位名", "original": "原文", "rewritten": "总领句\\n· 要点1\\n· 要点2", "note": "改写说明"}],
  "projects": [{"name": "项目名", "original": "原文", "rewritten": "总领句\\n· 要点1\\n· 要点2", "note": "改写说明"}],
  "skills": [{"before": "原技能描述", "after": "优化后描述"}],
  "actions": ["针对该候选人的具体改进建议1", "建议2", "建议3"]
}"""
    else:
        instruction = """你是一位专注于应届生求职的简历改写顾问，帮助候选人用真实、简洁的语言重新呈现经历，让HR一眼看到价值。

【改写原则 - 深度精修】
核心目标：每段经历都要让HR扫一眼总领句就能判断候选人的价值，然后通过2-3个要点了解具体做了什么。

【每段经历的结构要求（最重要）】
每段经历的 rewritten 必须包含：
1. 总领句（第一行，不加「· 」）：结果 / 交付物先行，行动在后
   HR筛简历扫第一句话，「完成了什么」比「负责了什么」更有说服力
   格式参考：「[完成了什么 / 交付了什么]，[独立程度 / 覆盖范围]」
   示例（好）：「独立完成B端供应链模块从0到1落地，功能覆盖财务/采购/行政3个业务方」
   示例（差）：「负责B端供应链模块的产品设计，推动从调研到上线全流程」← 动词先行，结果靠后
   注意：总领句不要重复具体细节，只说最重要的那一件事

2. 2-3条具体要点（以「· 」开头），用压缩STAR结构——完整STAR简历版面放不下，取「背景一句 + 行动 + 结果」精简版：
   第1条：背景（一句话，说明在解决什么问题/面对什么挑战）+ 核心行动
   示例：「· 针对旧版审批流程操作步骤繁琐导致投诉，主导需求调研（访谈15位供应商），定位3个核心卡点」
   第2条：具体行动 + 产出（方法/工具/交付物），动词开头，20-40字
   第3条（可选）：量化结果 或 具体交付物规模

【关于数据 — 铁律】
原文有数字 → 直接使用，不得修改幅度或单位
原文没有数字但该场景通常有量化产出 → 用占位符 [X%] 或 [X个] 标注，不捏造具体数值
  示例：「推动转化率提升 [X%]」「服务用户规模达 [X] 万」
  同时在对应经历的 note 字段注明：「建议候选人补充实际数据替换 [X%]」
原文没有数字且场景不适合量化 → 用具体动作范围替代
✓ 合法替代：「独立负责从调研到落地全流程」「覆盖3个核心业务方」
✗ 严禁捏造：直接写出「提升转化率30%」「用户满意度提升20%」等具体数字

【rewritten 字段格式】
总领句（不加·）\n· 要点1（背景+行动）\n· 要点2\n· 要点3（可选）

示例输出：
"独立完成B端供应商管理模块从0到1落地，覆盖采购/财务/行政3个业务部门\n· 针对旧系统审批流程繁琐导致供应商投诉，主导用户调研（访谈15人），定位3个核心流程卡点\n· 设计新版审批流程原型，减少4步操作，经历2轮评审迭代后上线"

【绝对禁止的词语和句式】
空洞强度词：深入、深度、充分、全面、系统性、多维度、有力、大力、有效（单独修饰时）、显著（无数字时）
企业黑话：赋能、闭环、打通、沉淀、对齐、拉通、助力、推动落地、全面负责、抓手、底层逻辑
自我评价：积极参与、深度参与、具有较强的、热爱、执行力强、自驱力强
虚假结果：取得良好成效、获得一致好评、受到广泛认可、产生积极影响（无具体表现时）
AI句式：「通过…实现了…」「基于…推动了…」「围绕…开展了…」「不断优化…」「持续推进…」「有力支撑…」
修饰语：不仅…还…更…、在…的背景下…

【note 字段格式】
一句话说明改写侧重点（如"总领句突出了独立负责的能力，与JD核心诉求匹配"）；若无JD则留空。

【输出格式（严格JSON）】
{
  "internships": [{"company": "公司名", "position": "岗位名", "original": "原文", "rewritten": "总领句\\n· 要点1\\n· 要点2", "note": "改写说明"}],
  "projects": [{"name": "项目名", "original": "原文", "rewritten": "总领句\\n· 要点1\\n· 要点2", "note": "改写说明"}],
  "skills": [{"before": "原技能描述", "after": "优化后描述"}],
  "actions": ["针对该候选人的具体改进建议1", "建议2", "建议3"]
}"""

    return f"""{instruction}{jd_section}

【简历原文】
{resume_text}

注意：
1. 必须识别并处理简历中的每一段实习经历和项目经历，不得遗漏
2. 每段经历都必须有总领句+2-3个要点的结构
3. 只输出JSON，不要输出任何解释文字"""

def build_parse_structure_prompt(resume_text: str) -> str:
    return f"""从以下简历中提取结构化信息，只提取原文真实存在的内容，不推断不补充。

【简历原文】
{resume_text}

【输出格式（严格JSON，不要输出任何其他文字）】
{{
  "internships": [
    {{
      "company": "公司名",
      "position": "岗位名",
      "department": "部门（原文没有则留空）",
      "start_date": "开始时间如2024.06（没有则留空）",
      "end_date": "结束时间（在职或至今则留空）",
      "is_current": false,
      "highlights": "原文经历描述，多条用\\n分隔，保留原文表述"
    }}
  ],
  "projects": [
    {{
      "name": "项目名",
      "role": "担任角色",
      "start_date": "开始时间（没有则留空）",
      "end_date": "结束时间（没有则留空）",
      "background": "项目背景（原文有则填，没有则留空）",
      "contribution": "主要贡献（原文有则填）",
      "result": "成果（原文有则填）",
      "tech_stack": "技术栈（原文有则填）"
    }}
  ],
  "skills": [
    {{
      "category": "分类（产品工具/数据分析/编程能力/软技能/行业知识 中选一）",
      "skill_name": "具体技能名称",
      "level": "熟练"
    }}
  ]
}}"""


router = APIRouter()


# ===== POST /api/v1/resume/parse-image — 截图粘贴OCR识别 =====
class ImageOCRRequest(BaseModel):
    image_base64: str   # base64 编码的图片（含 data:image/... 前缀）
    hint: str = "jd"    # "jd" 表示这是JD截图


@router.post("/parse-image")
async def parse_image_ocr(req: ImageOCRRequest):
    """
    用 LLM Vision API 识别截图中的文字
    支持 OpenAI/DeepSeek vision 兼容格式
    降级方案：若不支持 vision，返回错误让前端 fallback 到手动粘贴
    """
    import re as _re
    from core.config import settings

    if not _llm_key():
        raise HTTPException(status_code=400, detail="未配置 API Key，请先到设置页填写")

    # 确保 base64 有正确格式
    b64 = req.image_base64
    if not b64.startswith("data:"):
        b64 = "data:image/png;base64," + b64

    # 选择支持 vision 的模型
    vision_model_map = {
        "deepseek-chat": "deepseek-vl2",           # DeepSeek vision 模型
        "gpt-4o-mini": "gpt-4o-mini",              # OpenAI 原生支持
        "gpt-4o": "gpt-4o",
        "qwen-plus": "qwen-vl-plus",               # 通义千问 vision
    }
    base_model = settings.llm_model
    vision_model = vision_model_map.get(base_model, base_model)

    hint_text = {
        "jd": "这是一张职位招聘JD（岗位描述）的截图。请将截图中所有文字原样提取出来，保持段落结构，不要添加任何解释或评论。",
        "resume": "这是一张简历截图。请将截图中所有文字原样提取出来，保持格式结构，不要添加任何解释或评论。",
    }.get(req.hint, "请将截图中所有文字原样提取出来。")

    try:
        from openai import OpenAI
        client = OpenAI(api_key=_llm_key(), base_url=settings.llm_base_url)
        response = client.chat.completions.create(
            model=vision_model,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": hint_text},
                    {"type": "image_url", "image_url": {"url": b64}},
                ],
            }],
            max_tokens=2000,
        )
        text = response.choices[0].message.content or ""
        return {"text": text.strip(), "model_used": vision_model}

    except Exception as e:
        err_msg = str(e)
        # 针对常见错误给出友好提示
        if "vision" in err_msg.lower() or "image" in err_msg.lower() or "multimodal" in err_msg.lower():
            raise HTTPException(
                status_code=422,
                detail="当前配置的模型不支持图片识别（Vision），请手动将JD文字复制粘贴到输入框中"
            )
        raise HTTPException(status_code=500, detail=f"图片识别失败：{err_msg}")


# ===== POST /api/v1/resume/upload — 上传简历文件，返回提取的文本 =====
@router.post("/upload")
async def upload_resume(file: UploadFile = File(...)):
    """
    上传 PDF/TXT/DOCX 文件，返回提取的纯文本
    """
    import re as _re

    filename = file.filename or ""
    # 先读取文件内容（FastAPI async 环境里必须用 await）
    content = await file.read()

    try:
        name_lower = filename.lower()

        if name_lower.endswith(".pdf"):
            # PDF：用 PyPDF2 解析字节
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                texts = [p.extract_text() for p in reader.pages if p.extract_text()]
                text = "\n".join(texts).strip()
                if not text:
                    raise HTTPException(status_code=400, detail="PDF 文本为空，请粘贴文字版简历")
            except ImportError:
                raise HTTPException(status_code=400, detail="未安装 PyPDF2，请运行 pip install PyPDF2")

        elif name_lower.endswith(".docx"):
            # DOCX：必须用 python-docx，不能直接解码（DOCX 是 ZIP 二进制格式）
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                # 提取所有段落文字
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # 同时提取表格里的文字（简历常用表格布局）
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())
                text = "\n".join(paragraphs)
            except ImportError:
                raise HTTPException(
                    status_code=400,
                    detail="服务器未安装 python-docx，请将简历另存为 PDF 或 TXT 格式后重新上传"
                )
            except Exception as docx_err:
                raise HTTPException(status_code=400, detail=f"DOCX 解析失败：{docx_err}")

        else:
            # TXT 或其他：尝试多种编码
            text = None
            for enc in ("utf-8", "gbk", "gb2312", "utf-16"):
                try:
                    text = content.decode(enc)
                    break
                except Exception:
                    continue
            if text is None:
                text = content.decode("utf-8", errors="replace")

        if not text or not text.strip():
            raise HTTPException(status_code=400, detail="文件内容为空，请检查文件")

        return {
            "filename": filename,
            "text": text.strip(),
            "char_count": len(text.strip()),
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件解析失败：{e}")


# ===== 改写请求 Schema =====
class RewriteRequest(BaseModel):
    resume_text: str
    jd_text: str = ""
    target_role: str = "B端产品经理"
    mode: str = "light"   # "light"=轻修(temperature=0.3)  "deep"=精修(temperature=0.8)


# temperature 对应表
TEMPERATURE_MAP = {
    "light": 0.3,   # 轻修：保守，改动幅度小
    "deep":  0.6,   # 精修：有创意但不乱发挥，降低捏造风险
}

SYSTEM_PROMPT = """你是一位专注于应届生求职的简历改写顾问。核心原则：只基于候选人真实经历改写，绝不捏造数据，让简历读起来像候选人自己写的——具体、简洁、有主见。

【全局禁用词表】以下词语在任何情况下不得出现：

空洞强度词（假装表达程度，实际什么都没说）：
深入、深度、充分、全面、系统性、多维度、有力、大力、有效、显著、大幅、全方位

企业黑话（听起来高级但信息量为零）：
赋能、赋力、闭环、打通、拉通、对齐、沉淀、协同（单独作动词时）、推动落地、抓手、
颗粒度、底层逻辑、顶层设计、差异化（单独使用时）、链路

自我评价套话（声称而非证明能力）：
具有较强的、拥有丰富的、具备扎实的、热爱、自驱力强、执行力强、团队精神、勇于挑战

虚假结果词（没有数字支撑时禁用）：
显著提升（无数字）、大幅优化（无数字）、取得良好成效、获得一致好评、受到广泛认可、产生积极影响

AI特有句式（结构性AI味，必须拆解重写）：
通过…实现了…、基于…推动了…、围绕…开展了…、不断优化…、持续推进…、充分调研了…、有力支撑了…

【判断原则】如果一个词/句删掉之后意思没有变化，就不该出现。"""


def _parse_rewrite_json(raw: str) -> dict:
    """
    从 AI 返回的文字中提取 JSON
    处理 AI 可能加的 ```json ... ``` 包裹
    如果解析失败，返回降级格式（把原始文字放进 fallback 字段）
    """
    import json
    import re

    # 移除可能的 ```json ... ``` 包裹
    text = raw.strip()
    match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
    if match:
        text = match.group(1).strip()

    try:
        data = json.loads(text)
        # 验证必要字段存在
        return {
            "internships": data.get("internships", []),
            "projects":    data.get("projects", []),
            "skills":      data.get("skills", []),
            "actions":     data.get("actions", []),
            "raw":         None,  # 正常解析时不需要原文
        }
    except (json.JSONDecodeError, Exception):
        # 解析失败：降级返回原始文字，前端用文字模式展示
        return {
            "internships": [],
            "projects":    [],
            "skills":      [],
            "actions":     [],
            "raw": raw,   # 把原始 AI 输出给前端展示
        }


# ===== POST /api/v1/resume/rewrite — 简历改写（同步版，返回结构化 JSON）=====
@router.post("/rewrite")
async def rewrite_resume(req: RewriteRequest):
    """
    简历改写接口，返回结构化 JSON
    - mode=light  → temperature=0.3，轻修，保留风格只优化表达
    - mode=deep   → temperature=0.8，精修，STAR重写+JD匹配
    """
    from openai import OpenAI
    from core.config import settings

    if not _llm_key():
        raise HTTPException(status_code=400, detail="未配置 API Key，请在设置页配置")

    client = OpenAI(api_key=_llm_key(), base_url=settings.llm_base_url)
    temperature = TEMPERATURE_MAP.get(req.mode, 0.5)
    prompt = build_rewrite_prompt(req.resume_text, req.jd_text, req.mode)

    response = client.chat.completions.create(
        model=settings.llm_model,
        temperature=temperature,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
    )
    raw = response.choices[0].message.content
    parsed = _parse_rewrite_json(raw)

    return {
        **parsed,
        "mode":        req.mode,
        "temperature": temperature,
    }


# ===== POST /api/v1/resume/rewrite/stream — 简历改写（流式）=====
@router.post("/rewrite/stream")
async def rewrite_resume_stream(req: RewriteRequest):
    """
    流式简历改写，前端实时看到 AI 边写边输出
    支持 mode 参数控制改写力度
    """
    from openai import AsyncOpenAI
    from core.config import settings

    if not _llm_key():
        raise HTTPException(status_code=400, detail="未配置 API Key")

    client = AsyncOpenAI(api_key=_llm_key(), base_url=settings.llm_base_url)
    temperature = TEMPERATURE_MAP.get(req.mode, 0.5)
    prompt = build_rewrite_prompt(req.resume_text, req.jd_text, req.mode)

    async def generate():
        async with client.chat.completions.stream(
            model=settings.llm_model,
            temperature=temperature,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
        ) as stream:
            async for chunk in stream:
                if chunk.choices and chunk.choices[0].delta.content:
                    yield chunk.choices[0].delta.content

    return StreamingResponse(generate(), media_type="text/event-stream")


# ===== 保留旧接口（兼容首页）=====
class DiagnoseRequest(BaseModel):
    resume_text: str
    target_role: str = "B端产品经理"


@router.post("/diagnose")
async def diagnose_resume(req: DiagnoseRequest):
    """旧版诊断接口（精修模式，兼容首页调用）"""
    from openai import OpenAI
    from core.config import settings
    if not _llm_key():
        raise HTTPException(status_code=400, detail="未配置 API Key")
    client = OpenAI(api_key=_llm_key(), base_url=settings.llm_base_url)
    prompt = build_rewrite_prompt(req.resume_text, "", "deep")
    response = client.chat.completions.create(
        model=settings.llm_model,
        temperature=0.7,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
    )
    return {"result": response.choices[0].message.content}


@router.post("/diagnose/stream")
async def diagnose_resume_stream(req: DiagnoseRequest):
    """旧版流式诊断（精修模式）"""
    from openai import AsyncOpenAI
    from core.config import settings
    if not _llm_key():
        raise HTTPException(status_code=400, detail="未配置 API Key")
    client = AsyncOpenAI(api_key=_llm_key(), base_url=settings.llm_base_url)
    prompt = build_rewrite_prompt(req.resume_text, "", "deep")

    async def generate():
        async with client.chat.completions.stream(
            model=settings.llm_model, temperature=0.7,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
        ) as stream:
            async for chunk in stream:
                if chunk.choices and chunk.choices[0].delta.content:
                    yield chunk.choices[0].delta.content

    return StreamingResponse(generate(), media_type="text/event-stream")


# ===== GET /api/v1/resume/versions — 获取简历版本列表 =====
@router.get("/versions")
def get_resume_versions(application_id: Optional[int] = None):
    """获取版本库里的简历版本；传 application_id 则只返回该投递关联的版本"""
    with get_db() as db:
        if application_id is not None:
            rows = db.execute(
                "SELECT * FROM resume_versions WHERE application_id = ? ORDER BY created_at DESC",
                (application_id,),
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT * FROM resume_versions ORDER BY created_at DESC"
            ).fetchall()
        return [dict(row) for row in rows]


# ===== POST /api/v1/resume/versions — 保存新版本 =====
class VersionCreate(BaseModel):
    version_name: str
    content: str
    target_role: str = ""
    analysis_json: str = ""
    snapshot_json: str = ""
    application_id: Optional[int] = None


@router.post("/versions")
def create_resume_version(data: VersionCreate):
    """保存简历版本到版本库，可关联一条投递记录"""
    with get_db() as db:
        cursor = db.execute(
            "INSERT INTO resume_versions (version_name, content, target_role, analysis_json, snapshot_json, application_id) VALUES (?, ?, ?, ?, ?, ?)",
            (data.version_name, data.content, data.target_role, data.analysis_json, data.snapshot_json, data.application_id),
        )
        return {"id": cursor.lastrowid, "version_name": data.version_name}


# ===== PATCH /api/v1/resume/versions/{id}/link — 追溯绑定投递记录 =====
class VersionLink(BaseModel):
    application_id: Optional[int] = None


@router.patch("/versions/{version_id}/link")
def link_version_to_application(version_id: int, data: VersionLink):
    """将已有简历版本关联到某条投递记录（或解除关联）"""
    with get_db() as db:
        row = db.execute("SELECT id FROM resume_versions WHERE id = ?", (version_id,)).fetchone()
        if not row:
            from fastapi import HTTPException
            raise HTTPException(status_code=404, detail="版本不存在")
        db.execute(
            "UPDATE resume_versions SET application_id = ? WHERE id = ?",
            (data.application_id, version_id),
        )
        updated = db.execute("SELECT * FROM resume_versions WHERE id = ?", (version_id,)).fetchone()
        return dict(updated)


# ===== DELETE /api/v1/resume/versions/{id} — 删除版本 =====
@router.delete("/versions/{version_id}")
def delete_resume_version(version_id: int):
    """删除简历版本"""
    with get_db() as db:
        db.execute("DELETE FROM resume_versions WHERE id = ?", (version_id,))
        return {"message": f"已删除版本 {version_id}"}


# ===== POST /api/v1/resume/parse-and-sync — 解析简历并写入素材库 =====
class ParseSyncRequest(BaseModel):
    resume_text: str

@router.post("/parse-and-sync")
async def parse_and_sync(req: ParseSyncRequest):
    """
    用 AI 解析简历文本，提取实习/项目/技能结构，去重后写入素材库。
    返回新增数量，前端用于显示 toast。
    """
    import json, re
    from openai import OpenAI
    from core.config import settings

    if not _llm_key():
        raise HTTPException(status_code=400, detail="未配置 API Key")

    # 1. AI 解析
    client = OpenAI(api_key=_llm_key(), base_url=settings.llm_base_url)
    prompt = build_parse_structure_prompt(req.resume_text)
    response = client.chat.completions.create(
        model=settings.llm_model,
        temperature=0.1,
        messages=[
            {"role": "system", "content": "你是一个精准的简历信息提取助手，只输出JSON，不输出任何解释。"},
            {"role": "user", "content": prompt},
        ],
    )
    raw = response.choices[0].message.content or ""
    match = re.search(r'```(?:json)?\s*([\s\S]*?)```', raw)
    if match:
        raw = match.group(1).strip()
    try:
        parsed = json.loads(raw)
    except Exception:
        return {"added_internships": 0, "added_projects": 0, "added_skills": 0, "error": "解析失败"}

    interns  = parsed.get("internships", [])
    projects = parsed.get("projects", [])
    skills   = parsed.get("skills", [])

    added_i = added_p = added_s = 0

    with get_db() as db:
        # 获取已有数据用于去重
        existing_interns  = {(r["company"], r["position"]) for r in db.execute("SELECT company, position FROM profile_internship").fetchall()}
        existing_projects = {r["name"] for r in db.execute("SELECT name FROM profile_project").fetchall()}
        existing_skills   = {r["skill_name"] for r in db.execute("SELECT skill_name FROM profile_skill").fetchall()}

        for i in interns:
            key = (i.get("company", "").strip(), i.get("position", "").strip())
            if key[0] and key not in existing_interns:
                db.execute(
                    "INSERT INTO profile_internship (company, position, department, start_date, end_date, is_current, highlights) VALUES (?,?,?,?,?,?,?)",
                    (key[0], key[1], i.get("department",""), i.get("start_date",""),
                     i.get("end_date",""), 1 if i.get("is_current") else 0, i.get("highlights",""))
                )
                added_i += 1

        for p in projects:
            name = p.get("name", "").strip()
            if name and name not in existing_projects:
                db.execute(
                    "INSERT INTO profile_project (name, role, start_date, end_date, background, contribution, result, tech_stack) VALUES (?,?,?,?,?,?,?,?)",
                    (name, p.get("role",""), p.get("start_date",""), p.get("end_date",""),
                     p.get("background",""), p.get("contribution",""), p.get("result",""), p.get("tech_stack",""))
                )
                added_p += 1

        for s in skills:
            skill_name = s.get("skill_name", "").strip()
            if skill_name and skill_name not in existing_skills:
                db.execute(
                    "INSERT INTO profile_skill (category, skill_name, level) VALUES (?,?,?)",
                    (s.get("category","产品工具"), skill_name, s.get("level","熟练"))
                )
                added_s += 1

    return {
        "added_internships": added_i,
        "added_projects":    added_p,
        "added_skills":      added_s,
    }
